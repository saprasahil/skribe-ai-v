import streamlit as st
from openai import OpenAI
from docx import Document
import PyPDF2
import io
import os

try:
    from dotenv import load_dotenv  # type: ignore
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY") or st.secrets["OPENAI_API_KEY"]
except:
    api_key = st.secrets["OPENAI_API_KEY"]

# Secure API key (loaded from env or Streamlit secrets)
client = OpenAI(api_key=api_key)

def extract_text_from_file(uploaded_file):
    if uploaded_file.name.endswith('.pdf'):
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ''.join(page.extract_text() for page in reader.pages if page.extract_text())
    elif uploaded_file.name.endswith('.docx'):
        doc = Document(uploaded_file)
        text = '\n'.join([para.text for para in doc.paragraphs])
    elif uploaded_file.name.endswith('.txt'):
        text = uploaded_file.read().decode('utf-8')
    else:
        text = ""
    return text

def call_openai_gpt(prompt):
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error calling OpenAI: {e}"

def generate_cover_letter_gpt(job_desc, resume_text):
    prompt = f"""
Write a personalized cover letter for this job description and resume.
Make it 250–300 words in 4 paragraphs from which third paragraph should have 3 pointers in it. It should be professional, and enthusiastic.

Job Description:
{job_desc}

Resume:
{resume_text}
"""
    return call_openai_gpt(prompt)

def suggest_resume_improvements(job_desc, resume_text):
    prompt = f"""
You are an AI resume coach.

A job seeker has shared their resume and a job description. Your task is to review the resume and suggest improvements 
to make it more aligned to the job.

❌ Do not rewrite the resume.
✅ Provide 6 specific, actionable suggestions in plain English.

Resume:
{resume_text}

Job Description:
{job_desc}

Suggestions:
"""
    return call_openai_gpt(prompt)

def save_as_docx(text):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.set_page_config(page_title="Skribe AI", layout="centered")
st.title("📄 Skribe – AI Powered Job Application Assistant")

st.header("Step 1: Provide Job Description")
input_method = st.radio("Provide job description:", ["Paste Text", "Upload File"])
job_description = ""

if input_method == "Paste Text":
    job_description = st.text_area("Paste job description here:")
else:
    job_file = st.file_uploader("Upload job description (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])
    if job_file:
        job_description = extract_text_from_file(job_file)

st.header("Step 2: Upload Your Resume")
resume_file = st.file_uploader("Upload your resume (.pdf or .docx)", type=["pdf", "docx"])
resume_text = ""
if resume_file:
    resume_text = extract_text_from_file(resume_file)

if st.button("✨ Generate Cover Letter & Resume Suggestions"):
    if not job_description or not resume_text:
        st.error("Please provide both job description and resume.")
    else:
        with st.spinner("Generating with GPT-3.5..."):
            cover_letter = generate_cover_letter_gpt(job_description, resume_text)
            resume_suggestions = suggest_resume_improvements(job_description, resume_text)

        st.success("✅ Documents ready!")
        st.download_button("📥 Download Cover Letter (DOCX)", save_as_docx(cover_letter), "cover_letter.docx")

        st.subheader("💡 Suggestions to Personalise Your Resume")
        st.markdown(resume_suggestions)
